import pathlib
from pathlib import Path
import docx
import PyPDF2

class ScanDocument:  # catch the error if a file does not exist

    def __init__(self, src_dir) -> None:
        self.src_dir =  src_dir

    def get_docs(self) -> None:
        while True:
            self.dir_path = Path(self.src_dir)
            if self.dir_path.is_dir(): # If provided path exists in the OS.
                break
            else:
                self.src_dir = input("Invalid directory. Please enter a valid directory: ").strip()
                
        # Gets only the files in the root directory.
        self.present_files = [(file_path.name, file_path.suffix) for file_path in self.dir_path.iterdir() if file_path.is_file()]
        print("Files in Root Directory:")
        
        self.ext = {}
        for file_name, file_extension in self.present_files:
            print(f"{file_name}, ", end= "") # File names
            if file_extension in self.ext: # File extensions
                self.ext[file_extension] += 1
            else:
                self.ext[file_extension] = 1
        print("\n")
        print(f"Extensions:\n{self.ext}")
    
    def choose_file_type(self) -> str:
        self.pick_type = input("Which file type(s) do you want to scan? ").strip().lower()
        if self.pick_type == ".txt": # Read .txt files
            return self.read_txt()
        elif self.pick_type == ".docx":
            return self.read_docx()
        elif self.pick_type == ".pdf":
            return self.read_pdf()

    def read_txt(self):
        for file_name, file_ext in self.present_files:
                if file_ext == ".txt":
                    self.full_path = pathlib.PurePath(self.dir_path, file_name) # To avoid FileNotFoundError
                    with open(self.full_path, "r") as scanned_file: 
                        self.data = scanned_file.read()
                        return self.data

    def read_docx(self):
        self.lines = []
        for file_name, file_ext in self.present_files:
                if file_ext == ".docx":
                    self.full_path = pathlib.PurePath(self.dir_path, file_name)
                    doc = docx.Document(self.full_path) # Read .docx files
                    self.lines.extend([paragraph.text for paragraph in doc.paragraphs]) # Read all present .docx files 

                    if len(doc.tables) > 0:
                        for table in doc.tables:
                        # Access cells within the table
                            for row in table.rows:
                                for cell in row.cells:
                                    # Access content within each cell
                                    self.text = cell.text.replace('\n', '').strip()
                                    self.lines.append(self.text)
        return self.lines

    def read_pdf(self):
        self.pdf_info_list = []
        for file_name, file_ext in self.present_files:
                    
            self.pdf_info = {}

            if file_ext == ".pdf":
                self.full_path = pathlib.PurePath(self.dir_path, file_name)
                self.all_pdfs = str(self.full_path) # The conversion helps to access the pdf files easier
                with open(self.all_pdfs, "rb") as pdf_file_obj:
                    self.pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
                    for page_number in range(self.pdf_reader.getNumPages()):
                        self.page_obj = self.pdf_reader.getPage(page_number)
                        self.page_text = self.page_obj.extractText().replace("\n", "")
                        self.keys = f"Page {page_number + 1}"
                        self.pdf_info[self.keys] = self.page_text
                    self.pdf_info_list.append(self.pdf_info)  

        return self.pdf_info_list



def main():
    scan_document = ScanDocument(src_dir= input("Directory: ").strip())
    scan_document.get_docs()
    scan_document.choose_file_type()

if __name__ == "__main__":
    main()