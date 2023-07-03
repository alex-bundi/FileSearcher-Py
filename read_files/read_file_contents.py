import pathlib
from pathlib import Path
import docx
import PyPDF2

class ReadContents:
    """The class retrieves the information stored within a document based on the type of file extension it has."""
    present_files = []
    src_path = ""
        
    def read_txt(self) -> list:
        """Only reads documents with the file extension .txt"""
        self.txt_package = []
        self.txt_info = {}
        
        for file_name, file_ext in self.present_files:
            if file_ext == ".txt":
                self.full_path = pathlib.PurePath(self.src_path, file_name) # To avoid FileNotFoundError
                with open(self.full_path, "r") as scanned_file: 
                    self.data = scanned_file.read()
                    self.txt_info[file_name] = self.data
        self.txt_package.append(self.txt_info)

        return self.txt_package
        

    def read_docx(self) -> list:
        """Only reads documents with the file extension .docx"""

        self.lines = []
        for file_name, file_ext in self.present_files:
            if file_ext == ".docx":
                self.full_path = pathlib.PurePath(self.src_path, file_name)
                doc = docx.Document(self.full_path) # Read .docx files
                self.lines.extend([paragraph.text for paragraph in doc.paragraphs]) # Read all present .docx files 

                if len(doc.tables) > 0:
                    for table in doc.tables: # Access cells within the table
                        for row in table.rows:
                            for cell in row.cells:  # Access content within each cell
                                self.text = cell.text.replace('\n', '').strip()
                                self.lines.append(self.text)
        return self.lines

    def read_pdf(self) -> list:
        """Only reads documents with the file extension .pdf"""

        self.pdf_info_list = []
        for file_name, file_ext in self.present_files:       
            if file_ext == ".pdf":
                self.full_path = pathlib.PurePath(self.src_path, file_name)
                self.all_pdfs = str(self.full_path) # The conversion helps to access the pdf files easier

                with open(self.all_pdfs, "rb") as pdf_file_obj:
                    self.pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
                    for page_number in range(self.pdf_reader.getNumPages()):
                        self.pdf_info = {}
                        self.pdf_info["filename"] = self.full_path.name
                        self.page_obj = self.pdf_reader.getPage(page_number)
                        self.page_text = self.page_obj.extractText().replace("\n", "")
                        self.page_no = f"Page {page_number + 1}"
                        self.pdf_info["Page No"] = self.page_no
                        self.pdf_info["Content"] = self.page_text
                    self.pdf_info_list.append(self.pdf_info)  

        return self.pdf_info_list