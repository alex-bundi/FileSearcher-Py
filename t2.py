import pathlib
from pathlib import Path
import docx
import PyPDF2

class ScanDocument:  # catch the error if a file does not exist

    def __init__(self, parser) -> None:
        self.parser =  parser

    def get_docs(self):
        while True:
            self.src_dir = input("Directory: ")
            self.file_path = Path(self.src_dir)
            if self.file_path.is_dir() == True:
                break
            else:
                pass
        
        self.present_files = [self.file_path.name for self.file_path in self.file_path.iterdir() if self.file_path.is_file()]
        print(f"Files in Root Directory:\n{self.present_files}")

    def check_doc_type(self):
        """Gets the filename returns it and also the file type. """
        self.file_name = "issues.pdf"
        self.file_type = Path(self.file_name).suffix
        return self.file_name

    def scan_text(self):
        """Scans the text file depending on the file type."""
        self.check_doc_type()

        if self.file_type == ".txt": # Read .txt files
            with open(self.file_name, "r") as scanned_file: 
                self.data = scanned_file.read()
            return self.data
        elif self.file_type == ".docx":
            try:
                doc = docx.Document(self.file_name) # Read .docx files
                self.lines = " ".join([paragraph.text for paragraph in doc.paragraphs])
                return self.lines
            except docx.opc.exceptions.PackageNotFoundError: # If document does not exist
                print(f"Unable to scan opened document '{self.file_name}'.")
        elif self.file_type == ".pdf":
            with open(self.file_name, "rb") as pdf_file_obj:
                self.pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
                self.pdf_info = {}
                for page_number in range(self.pdf_reader.getNumPages()):
                    self.page_obj = self.pdf_reader.getPage(page_number)
                    self.page_text = self.page_obj.extractText().replace("\n", "")
                    self.keys = f"Page {page_number + 1}"
                    self.pdf_info[self.keys] = self.page_text
                return self.pdf_info
                
    def match_characters(self):
        self.check_doc_type()
        self.scan_text()

        if self.file_type == ".txt" or self.file_type == ".docx": # Full match
            if self.parser in self.scan_text():
                print()
            else:
                print("Phrase not found in the file.")
        elif self.file_type == ".pdf":
            for key, value in self.pdf_info.items():
                if self.parser in key or self.parser in value:
                    print("Phrase found in the file.")
                else:
                    print("not found")


def main():
    scan_document = ScanDocument(parser= input("Search for what: ").strip())
    scan_document.scan_text()
    scan_document.match_characters()

if __name__ == "__main__":
    main()